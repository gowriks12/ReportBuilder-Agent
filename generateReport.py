import docx
from docx.shared import Inches
import os
from utilFuncs import *
download_path = "C:\\Users\\gowri\\Documents\\Projects\\ReportBuilder-Agent\\Reports"
def createDoc(text):
    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading('Exploratory Data Analysis Report', 0)

    # Image in its native size
    doc.add_heading('General Information and Statistics', 3)
    doc.add_paragraph(text)
    doc.add_heading('Visualizations Generated', 3)

    images, recent_image = get_most_recent_image()

    for image in images:
        doc.add_picture(image, width=Inches(4), height=Inches(4))
        plots_folder = 'C:\\Users\\gowri\\Documents\\Projects\\ReportBuilder-Agent\\GeneratedPlots'
        move_image_to_folder(image, plots_folder)

    report_name = get_report_name()
    # Now save the document to a location

    doc.save(os.path.join(download_path, report_name))
    return os.path.join(download_path, report_name)

